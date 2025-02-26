from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_documentation():
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('ZMTech Finance - Stock Analysis Tool', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Introduction
        doc.add_paragraph('A Python-based desktop application for analyzing stock price movements and earnings impacts between pairs of stocks.')
        
        # Analysis Details Section
        doc.add_heading('Analysis Details', level=1)
        
        # Analysis Overview
        doc.add_heading('Analysis Overview', level=2)
        doc.add_paragraph('The tool performs comparative analysis between two stocks, focusing on:')
        overview_items = [
            'Stock price movements',
            'Earnings impact analysis',
            'Trading period comparisons before and after specific events'
        ]
        for item in overview_items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

        # Input Parameters
        doc.add_heading('Input Parameters', level=2)
        
        # Stock Selection
        doc.add_heading('Stock Selection', level=3)
        stock_params = [
            'Ticker 1: First stock symbol (e.g., AAPL, MSFT)',
            'Ticker 2: Second stock symbol for comparison'
        ]
        for param in stock_params:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(param)

        # Time Window Parameters
        doc.add_heading('Time Window Parameters', level=3)
        time_params = [
            'Days Before: Number of trading days to analyze before the event (default: 10)',
            'Days After: Number of trading days to analyze after the event (default: 10)'
        ]
        for param in time_params:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(param)

        # Analysis Components
        doc.add_heading('Analysis Components', level=2)
        doc.add_paragraph('The tool extracts and analyzes:')
        
        # Price Data
        doc.add_heading('Price Data', level=3)
        price_data = [
            'Historical price movements',
            'Daily price changes',
            'Trading volumes'
        ]
        for item in price_data:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

        # Comparative Analysis
        doc.add_heading('Comparative Analysis', level=3)
        comp_analysis = [
            'Relative price movements between the two stocks',
            'Correlation of price changes',
            'Impact of earnings announcements'
        ]
        for item in comp_analysis:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

        # Time Period Analysis
        doc.add_heading('Time Period Analysis', level=3)
        time_analysis = [
            'Pre-event price behavior',
            'Post-event price behavior',
            'Trading volume patterns'
        ]
        for item in time_analysis:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

        # Data Sources
        doc.add_heading('Data Sources', level=2)
        doc.add_paragraph('Stock data is retrieved using the yfinance (Yahoo Finance) API')
        doc.add_paragraph('Historical price data includes:')
        data_sources = [
            'Opening prices',
            'Closing prices',
            'High/Low prices',
            'Trading volumes',
            'Adjusted close prices'
        ]
        for source in data_sources:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(source)

        # Output Analysis
        doc.add_heading('Output Analysis', level=2)
        doc.add_paragraph('The tool provides:')
        
        # Numerical Results
        doc.add_heading('Numerical Results', level=3)
        num_results = [
            'Price change percentages',
            'Correlation coefficients',
            'Volume analysis'
        ]
        for result in num_results:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(result)

        # Statistical Measures
        doc.add_heading('Statistical Measures', level=3)
        stat_measures = [
            'Price movement patterns',
            'Trading volume patterns',
            'Relative performance metrics'
        ]
        for measure in stat_measures:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(measure)

        # Add note about implementation
        note = doc.add_paragraph()
        note.add_run('Note: ').bold = True
        note.add_run('For detailed implementation of the analysis logic, refer to the analyze_earnings_impact() function in the az.py module.')

        # Save the document
        output_path = Path(__file__).parent / 'ZMTech_Documentation.docx'
        doc.save(str(output_path))
        print(f"Documentation created successfully at: {output_path}")

    except Exception as e:
        print(f"Error creating documentation: {str(e)}")

if __name__ == '__main__':
    create_documentation()