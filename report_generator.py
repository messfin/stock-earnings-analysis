
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import pandas as pd
import datetime
import io

class ZMTechReport:
    def __init__(self, output_dir=None):
        self.output_dir = Path(output_dir) if output_dir else Path('zmtech_finance/reports')
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.setup_styles()
        
    def setup_styles(self):
        """Setup report styling"""
        self.colors = {
            'primary': RGBColor(30, 61, 89),    # #1E3D59
            'secondary': RGBColor(23, 184, 144), # #17B890
            'accent': RGBColor(255, 180, 0),     # #FFB400
            'text': RGBColor(51, 51, 51)         # #333333
        }
        
        self.styles = {
            'title': {
                'font': 'Montserrat',
                'size': Pt(24),
                'color': self.colors['primary']
            },
            'heading1': {
                'font': 'Montserrat',
                'size': Pt(18),
                'color': self.colors['primary']
            },
            'heading2': {
                'font': 'Montserrat',
                'size': Pt(14),
                'color': self.colors['secondary']
            },
            'body': {
                'font': 'Open Sans',
                'size': Pt(11),
                'color': self.colors['text']
            }
        }
        
    def generate_report(self, analysis_results, ticker1, ticker2):
        """Generate comprehensive analysis report"""
        doc = Document()
        self.add_header(doc, ticker1, ticker2)
        
        # Executive Summary
        self.add_executive_summary(doc, analysis_results, ticker1, ticker2)
        
        # Technical Analysis
        self.add_technical_analysis(doc, analysis_results)
        
        # Earnings Analysis
        self.add_earnings_analysis(doc, analysis_results)
        
        # Correlation Analysis
        self.add_correlation_analysis(doc, analysis_results)
        
        # Charts and Visualizations
        self.add_visualizations(doc, analysis_results)
        
        # Save report
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = self.output_dir / f'ZMTech_Analysis_{ticker1}_{ticker2}_{timestamp}.docx'
        doc.save(filename)
        return filename
        
    def add_header(self, doc, ticker1, ticker2):
        """Add report header"""
        # Title
        title = doc.add_heading(f'Stock Analysis Report: {ticker1} vs {ticker2}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(
            f'Generated on {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
        self.apply_style(date_run, 'body')
        
        doc.add_paragraph()  # Spacing
        
    def add_executive_summary(self, doc, results, ticker1, ticker2):
        """Add executive summary"""
        doc.add_heading('Executive Summary', level=1)
        
        summary = doc.add_paragraph()
        summary_text = f"""
        This report provides a comprehensive analysis comparing {ticker1} and {ticker2}. 
        Key findings include:
        """
        summary_run = summary.add_run(summary_text)
        self.apply_style(summary_run, 'body')
        
        # Add key metrics
        self.add_key_metrics_table(doc, results, ticker1, ticker2)
        
    def add_technical_analysis(self, doc, results):
        """Add technical analysis section"""
        doc.add_heading('Technical Analysis', level=1)
        
        tech_data = results['technical']
        
        for stock_name, analysis in tech_data.items():
            doc.add_heading(f'{stock_name} Technical Indicators', level=2)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add headers
            headers = table.rows[0].cells
            headers[0].text = 'Indicator'
            headers[1].text = 'Value'
            
            # Add data
            for indicator, value in analysis.items():
                row = table.add_row().cells
                row[0].text = indicator.replace('_', ' ').title()
                row[1].text = str(value)
                
    def add_earnings_analysis(self, doc, results):
        """Add earnings analysis section"""
        doc.add_heading('Earnings Analysis', level=1)
        
        earnings_data = results['earnings']
        
        for stock_name, earnings in earnings_data.items():
            if earnings is not None and not earnings.empty:
                doc.add_heading(f'{stock_name} Earnings History', level=2)
                
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Headers
                headers = table.rows[0].cells
                headers[0].text = 'Date'
                headers[1].text = 'Reported EPS'
                headers[2].text = 'Estimated EPS'
                headers[3].text = 'Surprise %'
                
                # Data
                for index, row in earnings.iterrows():
                    table_row = table.add_row().cells
                    table_row[0].text = index.strftime('%Y-%m-%d')
                    table_row[1].text = f"{row['EPS Actual']:.2f}"
                    table_row[2].text = f"{row['EPS Estimate']:.2f}"
                    table_row[3].text = f"{row['Surprise(%)']:.2f}%"
                    
    def add_correlation_analysis(self, doc, results):
        """Add correlation analysis section"""
        doc.add_heading('Correlation Analysis', level=1)
        
        corr_data = results['correlation']
        
        paragraph = doc.add_paragraph()
        corr_text = f"""
        The correlation coefficient between the two stocks is {corr_data['correlation']:.2f}.
        This indicates a {'strong' if abs(corr_data['correlation']) > 0.7 else 'moderate' if abs(corr_data['correlation']) > 0.3 else 'weak'} 
        {'positive' if corr_data['correlation'] > 0 else 'negative'} relationship.
        """
        corr_run = paragraph.add_run(corr_text)
        self.apply_style(corr_run, 'body')
        
    def add_visualizations(self, doc, results):
        """Add charts and visualizations"""
        doc.add_heading('Technical Charts', level=1)
        
        # Add each chart from results
        for chart_name, fig in results['charts'].items():
            # Save figure to memory
            img_stream = io.BytesIO()
            fig.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
            img_stream.seek(0)
            
            # Add to document
            doc.add_picture(img_stream, width=Inches(6))
            
            # Add caption
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(chart_name.replace('_', ' ').title())
            self.apply_style(caption_run, 'body')
            
    def apply_style(self, run, style_name):
        """Apply text style to a run"""
        style = self.styles[style_name]
        font = run.font
        font.name = style['font']
        font.size = style['size']
        font.color.rgb = style['color']
        
    def add_key_metrics_table(self, doc, results, ticker1, ticker2):
        """Add key metrics comparison table"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'Metric'
        headers[1].text = ticker1
        headers[2].text = ticker2
        
        # Add metrics
        metrics = [
            ('Current Price', 'Close'),
            ('RSI', 'RSI'),
            ('Trend', 'trend'),
            ('Above MA200', 'above_ma200'),
            ('Correlation', 'correlation')
        ]
        
        for metric_name, metric_key in metrics:
            row = table.add_row().cells
            row[0].text = metric_name
            row[1].text = str(results['technical']['stock1'].get(metric_key, 'N/A'))
            row[2].text = str(results['technical']['stock2'].get(metric_key, 'N/A'))

def main():
    # Test report generation
    from analysis_engine import ZMTechAnalysis
    
    analyzer = ZMTechAnalysis()
    results = analyzer.analyze_stocks('AAPL', 'MSFT')
    
    report_gen = ZMTechReport()
    report_file = report_gen.generate_report(results, 'AAPL', 'MSFT')
    print(f"Report generated: {report_file}")

if __name__ == "__main__":
    main()