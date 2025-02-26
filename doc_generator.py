
1. First, let's create a documentation generator (`doc_generator.py`):

```python
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import matplotlib.pyplot as plt
import datetime

class ZMTechDocument:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.assets_path = Path('zmtech_assets')
        
    def setup_styles(self):
        """Setup ZMTech document styles"""
        # Colors
        self.colors = {
            'primary_blue': RGBColor(30, 61, 89),  # #1E3D59
            'teal': RGBColor(23, 184, 144),        # #17B890
            'gold': RGBColor(255, 180, 0),         # #FFB400
            'platinum': RGBColor(232, 232, 232)    # #E8E8E8
        }
        
        # Heading Styles
        for level in range(1, 4):
            style = self.doc.styles.add_style(f'ZMTech Heading {level}', 
                                            WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Montserrat'
            font.size = Pt(20 - (level * 2))
            font.bold = True
            font.color.rgb = self.colors['primary_blue']
        
        # Body Style
        style = self.doc.styles.add_style('ZMTech Body', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Open Sans'
        font.size = Pt(11)
        
    def add_header(self):
        """Add ZMTech branded header"""
        header = self.doc.sections[0].header
        htable = header.add_table(1, 2, Inches(6))
        
        # Add logo
        logo_cell = htable.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(self.assets_path / 'zmtech_logo.png'), 
                           width=Inches(2))
        
        # Add date
        date_cell = htable.cell(0, 1)
        date_paragraph = date_cell.paragraphs[0]
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(
            datetime.datetime.now().strftime('%Y-%m-%d'))
        
    def create_stock_analysis_report(self, ticker1, ticker2, results):
        """Create stock analysis report"""
        # Add header
        self.add_header()
        
        # Title
        title = self.doc.add_paragraph(
            f'Stock Analysis Report: {ticker1} vs {ticker2}',
            style='ZMTech Heading 1')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        self.doc.add_heading('Executive Summary', level=2)
        summary = self.doc.add_paragraph(style='ZMTech Body')
        summary.add_run(
            f'Comparative analysis of {ticker1} and {ticker2} over the last 10 earnings periods.')
        
        # Technical Analysis
        self.doc.add_heading('Technical Analysis', level=2)
        
        # Add charts
        self.add_technical_charts(results)
        
        # Earnings Analysis
        self.doc.add_heading('Earnings Analysis', level=2)
        
        # Add tables
        self.add_earnings_tables(results)
        
        # Save document
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'ZMTech_Analysis_{ticker1}_{ticker2}_{timestamp}.docx'
        self.doc.save(filename)
        return filename
        
    def add_technical_charts(self, results):
        """Add technical analysis charts"""
        # Create charts using matplotlib
        plt.style.use('seaborn')
        
        # RSI Distribution
        fig, ax = plt.subplots(figsize=(10, 6))
        # ... chart creation code ...
        
        # Save and add to document
        chart_path = self.assets_path / 'temp_chart.png'
        plt.savefig(chart_path)
        self.doc.add_picture(str(chart_path))
        
    def add_earnings_tables(self, results):
        """Add earnings analysis tables"""
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add headers and data
        # ... table population code ...

def main():
    # Create document generator
    doc_gen = ZMTechDocument()
    
    # Example usage
    results = {
        'technical_data': {},
        'earnings_data': {}
    }
    
    doc_gen.create_stock_analysis_report('AAPL', 'MSFT', results)

if __name__ == "__main__":
    main()
```
