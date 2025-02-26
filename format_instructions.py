from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_instruction_doc(filename="setup_instructions.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading('Stock Analysis Setup Instructions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph('Follow these steps to set up and run the stock analysis scripts.')
    
    # Step 1: Clone Repository
    heading = doc.add_heading('1. Clone Repository', level=1)
    p = doc.add_paragraph()
    p.add_run('Execute the following commands in your terminal:').bold = True
    code = doc.add_paragraph('git clone https://github.com/messfin/stock-earnings-analysis.git\ncd stock-earnings-analysis')
    code.style = 'Code'
    
    # Step 2: Set Up Python Environment
    heading = doc.add_heading('2. Set Up Python Environment', level=1)
    p = doc.add_paragraph()
    p.add_run('Create and activate virtual environment:').bold = True
    
    # Windows instructions
    doc.add_paragraph('For Windows:')
    code = doc.add_paragraph('python -m venv venv\nvenv\\Scripts\\activate')
    code.style = 'Code'
    
    # Mac/Linux instructions
    doc.add_paragraph('For Mac/Linux:')
    code = doc.add_paragraph('python -m venv venv\nsource venv/bin/activate')
    code.style = 'Code'
    
    # Install packages
    p = doc.add_paragraph()
    p.add_run('Install required packages:').bold = True
    code = doc.add_paragraph('pip install pandas python-docx yfinance matplotlib')
    code.style = 'Code'
    
    # Step 3: Run Scripts
    heading = doc.add_heading('3. Run Analysis Scripts', level=1)
    p = doc.add_paragraph()
    p.add_run('Execute either of the following commands:').bold = True
    code = doc.add_paragraph('python earnings_sector_compare.py\n# or\npython zmtech_main.py')
    code.style = 'Code'
    
    # Step 4: Find Output
    heading = doc.add_heading('4. Locate Output', level=1)
    p = doc.add_paragraph('The analysis results will be saved as "stock_analysis_report.docx" in your current directory.')
    
    # Notes section
    doc.add_heading('Important Notes:', level=1)
    notes = doc.add_paragraph()
    notes.add_run('• Ensure you have Python installed on your system\n')
    notes.add_run('• Make sure you have adequate permissions in the directory\n')
    notes.add_run('• Check internet connectivity for stock data retrieval')
    
    # Save the document
    doc.save(filename)
    print(f"Instructions saved as {filename}")

if __name__ == "__main__":
    create_instruction_doc()