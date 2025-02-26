import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_master_doc():
    doc = Document()
    
    # Add styles
    styles = doc.styles
    style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Courier New'
    style.font.size = Pt(10)
    
    # Title Page
    doc.add_heading('Stock Analysis Project', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run('Complete Documentation').bold = True
    doc.add_paragraph(f'Generated: {os.path.basename(__file__)}')
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    sections = [
        "1. Project Setup",
        "2. GitHub Configuration",
        "3. Installation Guide",
        "4. Usage Instructions",
        "5. Script Documentation",
        "6. Troubleshooting",
        "7. API Reference"
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Number')
    doc.add_page_break()
    
    # 1. Project Setup
    doc.add_heading('1. Project Setup', level=1)
    add_project_setup_section(doc)
    doc.add_page_break()
    
    # 2. GitHub Configuration
    doc.add_heading('2. GitHub Configuration', level=1)
    add_github_config_section(doc)
    doc.add_page_break()
    
    # 3. Installation Guide
    doc.add_heading('3. Installation Guide', level=1)
    add_installation_section(doc)
    doc.add_page_break()
    
    # 4. Usage Instructions
    doc.add_heading('4. Usage Instructions', level=1)
    add_usage_section(doc)
    doc.add_page_break()
    
    # 5. Script Documentation
    doc.add_heading('5. Script Documentation', level=1)
    add_script_documentation(doc)
    doc.add_page_break()
    
    # 6. Troubleshooting
    doc.add_heading('6. Troubleshooting', level=1)
    add_troubleshooting_section(doc)
    doc.add_page_break()
    
    # 7. API Reference
    doc.add_heading('7. API Reference', level=1)
    add_api_reference(doc)
    
    # Save the document
    try:
        # Create docs directory if it doesn't exist
        os.makedirs('docs', exist_ok=True)
        doc.save('docs/master_documentation.docx')
        print("Master documentation created successfully in docs/master_documentation.docx")
    except Exception as e:
        desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
        doc.save(os.path.join(desktop_path, 'master_documentation.docx'))
        print(f"Documentation saved to Desktop due to error: {str(e)}")

def add_project_setup_section(doc):
    doc.add_paragraph("Follow these steps to set up the project environment:", style='Intense Quote')
    
    steps = [
        ("Create Project Directory", """
mkdir D:\\code\\stock_analysis
cd D:\\code\\stock_analysis"""),
        
        ("Clone Repository", """
git clone https://github.com/messfin/stock-earnings-analysis.git
cd stock-earnings-analysis"""),
        
        ("Create Virtual Environment", """
python -m venv venv
venv\\Scripts\\activate  # Windows
source venv/bin/activate  # Mac/Linux""")
    ]
    
    for title, code in steps:
        doc.add_heading(title, level=2)
        doc.add_paragraph(code, style='CodeBlock')

def add_github_config_section(doc):
    doc.add_paragraph("GitHub Configuration Steps:", style='Intense Quote')
    
    steps = [
        "1. Create GitHub Account",
        "2. Generate Personal Access Token (PAT)",
        "3. Configure Git Credentials",
        "4. Test Connection"
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph("""
git config --global user.name "YOUR_USERNAME"
git config --global user.email "YOUR_EMAIL"
git config --global credential.helper store""", style='CodeBlock')

def add_installation_section(doc):
    doc.add_paragraph("Required Package Installation:", style='Intense Quote')
    
    doc.add_paragraph("""
pip install pandas
pip install python-docx
pip install yfinance
pip install matplotlib""", style='CodeBlock')

def add_usage_section(doc):
    doc.add_heading("Running the Scripts", level=2)
    doc.add_paragraph("""
# Run earnings analysis
python earnings_sector_compare.py

# Run ZMTech analysis
python zmtech_main.py""", style='CodeBlock')

def add_script_documentation(doc):
    scripts = {
        "earnings_sector_compare.py": "Analyzes earnings across different market sectors",
        "zmtech_main.py": "Performs ZMTech specific analysis",
        "export.py": "Handles document generation and export functions"
    }
    
    for script, description in scripts.items():
        doc.add_heading(script, level=2)
        doc.add_paragraph(description)

def add_troubleshooting_section(doc):
    issues = [
        ("Permission Denied", "Run as administrator or check file permissions"),
        ("Import Errors", "Verify all required packages are installed"),
        ("GitHub Authentication", "Check PAT token and credentials"),
        ("Data Access Issues", "Verify internet connection and API access")
    ]
    
    for issue, solution in issues:
        doc.add_heading(issue, level=2)
        doc.add_paragraph(solution)

def add_api_reference(doc):
    apis = [
        ("YFinance API", "Stock data retrieval"),
        ("Pandas API", "Data manipulation"),
        ("Matplotlib API", "Data visualization")
    ]
    
    for api, description in apis:
        doc.add_heading(api, level=2)
        doc.add_paragraph(description)

if __name__ == "__main__":
    print("Generating master documentation...")
    create_master_doc()