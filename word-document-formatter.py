from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_formatted_document():
    # Create a new Document
    doc = Document()
    
    # Set margins for one-page format
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add title
    title = doc.add_heading('Steps to Clone and Run the Files', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Main steps
    doc.add_heading('1. Clone the Repository', level=2)
    p = doc.add_paragraph()
    p.add_run('git clone [repository-url]')
    
    doc.add_heading('2. Set Up Python Environment', level=2)
    p = doc.add_paragraph()
    p.add_run('cd [repository-directory]\n')
    p.add_run('pip install -r requirements.txt')
    
    doc.add_heading('3. Run the Files', level=2)
    p = doc.add_paragraph()
    p.add_run('python zmtech_main.py')
    
    # Important notes section
    doc.add_heading('Important Notes:', level=2)
    
    # Create bulleted list
    bullet_points = [
        "The repository contains a requirements.txt file which will install all necessary dependencies",
        "The project includes various analysis tools including technical indicators (RSI, Moving Averages), volume analysis, correlation studies, and interactive GUI for visualization",
        "Make sure you have Python installed on your system before starting",
        "If you encounter any issues, the repository includes documentation files that might help: technical_guide.md, troubleshooting.md, user_manual.md"
    ]
    
    for point in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(point)
    
    # Add section for Word document export
    doc.add_heading('Export to Word Document', level=2)
    
    p = doc.add_paragraph()
    p.add_run("To automatically export results to a Word document, add this code to your script:").bold = True
    
    # Code block
    code = """
# Install required package if needed: pip install python-docx
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

def save_to_word(data, filename="analysis_report.docx"):
    doc = Document()
    doc.add_heading('Analysis Report', 0)
    
    # Add data tables
    doc.add_heading('Data Summary', level=1)
    doc.add_paragraph(str(data))
    
    # Add any saved plots
    try:
        doc.add_heading('Visualizations', level=1)
        doc.add_picture('plot.png', width=Inches(6))
    except:
        pass
        
    doc.save(filename)
    print(f"Report saved as {filename}")

# Example usage
# save_to_word(your_data_variable)
"""
    
    p = doc.add_paragraph()
    p.add_run(code)
    
    # Save the document
    output_filename = "setup_instructions.docx"
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")

if __name__ == "__main__":
    create_formatted_document()
