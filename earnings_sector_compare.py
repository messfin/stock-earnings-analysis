import yfinance as yf
import pandas as pd
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import os
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches

class StockAnalyzer:
    def __init__(self):
        self.cache = {}

    def get_earnings_dates(self, ticker: str) -> List[datetime]:
        """Fetch historical earnings dates for a ticker"""
        try:
            stock = yf.Ticker(ticker)
            earnings_history = stock.earnings_dates
            
            if earnings_history is None or earnings_history.empty:
                return []
                
            # Ensure consistent timezone handling
            dates = earnings_history.index
            if dates.tz is not None:
                dates = dates.tz_localize(None)  # Remove timezone info
            
            return sorted(dates, reverse=True)[:8]
            
        except Exception as e:
            print(f"Error fetching earnings dates for {ticker}: {e}")
            return []
            
    def get_stock_data(self, ticker: str, start_date: datetime, end_date: datetime) -> Optional[pd.DataFrame]:
        """Get stock price data with caching"""
        try:
            # Ensure dates are timezone-naive
            start_date = pd.to_datetime(start_date).tz_localize(None)
            # Get more historical data for accurate MA calculations
            extended_start = start_date - timedelta(days=400)  # Get extra days for MA calculation
            end_date = pd.to_datetime(end_date).tz_localize(None)
            
            cache_key = f"{ticker}_{start_date}_{end_date}"
            if cache_key in self.cache:
                return self.cache[cache_key]
                
            stock = yf.Ticker(ticker)
            data = stock.history(start=extended_start, end=end_date)
            
            if not data.empty:
                # Ensure index is timezone-naive
                if data.index.tz is not None:
                    data.index = data.index.tz_localize(None)
                
                # Calculate returns
                data['Daily_Return'] = data['Close'].pct_change()
                data['Cumulative_Return'] = (1 + data['Daily_Return']).cumprod() - 1
                
                # Calculate RSI (14-day)
                delta = data['Close'].diff()
                gain = delta.where(delta > 0, 0)
                loss = -delta.where(delta < 0, 0)
                avg_gain = gain.rolling(window=14).mean()
                avg_loss = loss.rolling(window=14).mean()
                rs = avg_gain / avg_loss
                data['RSI'] = 100 - (100 / (1 + rs))
                
                # Calculate Moving Averages
                data['MA50'] = data['Close'].rolling(window=50).mean()
                data['MA200'] = data['Close'].rolling(window=200).mean()
                
                # Get IV (if available)
                try:
                    options = stock.options
                    if options:
                        nearest_option = stock.option_chain(options[0])
                        data['IV'] = nearest_option.calls['impliedVolatility'].mean()
                    else:
                        data['IV'] = None
                except:
                    data['IV'] = None
                
                # Trim the data back to the requested date range
                data = data[data.index >= start_date]
                
                self.cache[cache_key] = data
                return data
                
            return None
            
        except Exception as e:
            print(f"Error fetching data for {ticker}: {e}")
            return None

    def check_ma_signals(self, data: pd.DataFrame) -> Dict[str, str]:
        """Check moving average signals"""
        latest = data.iloc[-1]
        signals = {}
        
        # Check if price is above/below MAs
        if pd.notnull(latest['MA50']):
            signals['above_50ma'] = latest['Close'] > latest['MA50']
        if pd.notnull(latest['MA200']):
            signals['above_200ma'] = latest['Close'] > latest['MA200']
            
        # Check MA crossover
        if pd.notnull(latest['MA50']) and pd.notnull(latest['MA200']):
            signals['golden_cross'] = latest['MA50'] > latest['MA200']
            
        return signals

    def get_current_price(self, ticker: str) -> float:
        """Get real-time current price for a ticker"""
        try:
            stock = yf.Ticker(ticker)
            current = stock.history(period='1d')
            if not current.empty:
                return current['Close'].iloc[-1]
            return None
        except Exception as e:
            print(f"Error fetching current price for {ticker}: {e}")
            return None

    def get_current_iv(self, ticker: str) -> float:
        """Get current IV from the nearest expiration options"""
        try:
            stock = yf.Ticker(ticker)
            options = stock.options
            if options:
                nearest_option = stock.option_chain(options[0])
                # Average IV from both calls and puts
                call_iv = nearest_option.calls['impliedVolatility'].mean()
                put_iv = nearest_option.puts['impliedVolatility'].mean()
                return (call_iv + put_iv) / 2
            return None
        except Exception as e:
            print(f"Error fetching IV for {ticker}: {e}")
            return None

    def calculate_correlation(self, data1: pd.Series, data2: pd.Series) -> float:
        """Calculate correlation between two price series"""
        try:
            if data1 is None or data2 is None or len(data1) < 2 or len(data2) < 2:
                return None
                
            # Convert to returns for better correlation analysis
            returns1 = data1.pct_change().dropna()
            returns2 = data2.pct_change().dropna()
            
            # Align the data and calculate correlation
            aligned_returns = pd.concat([returns1, returns2], axis=1).dropna()
            if len(aligned_returns) < 2:  # Need at least 2 points for correlation
                return None
                
            correlation = aligned_returns.corr().iloc[0, 1]
            return correlation if not pd.isna(correlation) else None
            
        except Exception as e:
            print(f"Error calculating correlation: {e}")
            return None

    def get_correlation_category(self, corr: float) -> str:
        """Categorize correlation strength"""
        if corr is None:
            return "N/A"
        abs_corr = abs(corr)
        if abs_corr >= 0.7:
            return f"High ({corr:.2%})"
        elif abs_corr >= 0.4:
            return f"Moderate ({corr:.2%})"
        else:
            return f"Low ({corr:.2%})"

    def get_historical_iv(self, ticker: str, date: datetime) -> float:
        """Get historical IV for a specific date"""
        try:
            stock = yf.Ticker(ticker)
            # Get options expiring after the target date
            all_options = stock.options
            if not all_options:
                return None
                
            # Convert date to datetime and ensure it's timezone-naive
            target_date = pd.to_datetime(date).tz_localize(None)
            
            # Find the nearest expiration after the target date
            valid_dates = [pd.to_datetime(d).tz_localize(None) for d in all_options]
            future_dates = [d for d in valid_dates if d > target_date]
            if not future_dates:
                return None
                
            nearest_expiry = min(future_dates)
            
            # Get the option chain for that expiration
            option_chain = stock.option_chain(nearest_expiry.strftime('%Y-%m-%d'))
            if option_chain is None:
                return None
                
            # Get ATM options for more accurate IV
            current_price = self.get_current_price(ticker)
            if current_price is None:
                return None
                
            # Filter for near-the-money options
            calls = option_chain.calls
            puts = option_chain.puts
            
            # Get options closest to current price
            calls = calls[abs(calls['strike'] - current_price) < current_price * 0.1]  # Within 10% of current price
            puts = puts[abs(puts['strike'] - current_price) < current_price * 0.1]
            
            if calls.empty and puts.empty:
                return None
                
            # Calculate weighted average IV
            call_iv = calls['impliedVolatility'].mean() if not calls.empty else 0
            put_iv = puts['impliedVolatility'].mean() if not puts.empty else 0
            
            count = (0 if calls.empty else 1) + (0 if puts.empty else 1)
            if count == 0:
                return None
                
            return (call_iv + put_iv) / count
            
        except Exception as e:
            print(f"Error fetching historical IV for {ticker} at {date}: {e}")
            return None

class ERAnalysisApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Earnings Analysis")
        self.analyzer = StockAnalyzer()
        
        # Create frames
        self.input_frame = ttk.Frame(self.root, padding="10")
        self.input_frame.grid(row=0, column=0, sticky="nsew")
        
        # Add column control frame
        self.column_frame = ttk.Frame(self.root, padding="10")
        self.column_frame.grid(row=1, column=0, sticky="nsew")
        
        self.summary_frame = ttk.Frame(self.root, padding="10")
        self.summary_frame.grid(row=2, column=0, sticky="nsew")
        
        self.charts_frame = ttk.Frame(self.root, padding="10")
        self.charts_frame.grid(row=3, column=0, sticky="nsew")
        
        self.export_frame = ttk.Frame(self.root, padding="10")
        self.export_frame.grid(row=4, column=0, sticky="nsew")
        
        # Define all possible columns
        self.all_columns = {
            'Ticker': True,
            'Current Price': True,
            'Pre-ER Price': True,
            'Post-ER Price': True,
            'Price Change': True,
            'Current IV': True,
            'Pre-ER IV': True,
            'Post-ER IV': True,
            'IV Change': True,
            'Pre-ER Return': True,
            'Post-ER Return': True,
            'Volume Change': True,
            'Current RSI': True,
            'Price vs MA200': True,
            'Price vs MA50': True,
            'MA Cross': True,
            'Correlation': True
        }
        
        # Create column toggles
        self.create_column_controls()
        
        # Create input widgets
        ttk.Label(self.input_frame, text="Ticker:").grid(row=0, column=0)
        self.ticker_entry = ttk.Entry(self.input_frame)
        self.ticker_entry.grid(row=0, column=1)
        self.ticker_entry.bind('<FocusOut>', self.populate_earnings_dates)
        
        ttk.Label(self.input_frame, text="Earnings Date:").grid(row=1, column=0)
        self.er_date_var = tk.StringVar()
        self.er_date_combo = ttk.Combobox(self.input_frame, textvariable=self.er_date_var, state='readonly')
        self.er_date_combo.grid(row=1, column=1)
        
        ttk.Label(self.input_frame, text="Days Around ER:").grid(row=2, column=0)
        self.days_entry = ttk.Entry(self.input_frame)
        self.days_entry.insert(0, "5")
        self.days_entry.grid(row=2, column=1)
        
        ttk.Label(self.input_frame, text="Peers (comma-separated):").grid(row=3, column=0)
        self.peers_entry = ttk.Entry(self.input_frame)
        self.peers_entry.grid(row=3, column=1)
        
        ttk.Button(self.input_frame, text="Analyze", command=self.run_analysis).grid(row=4, column=0, columnspan=2)
        
        # Add export buttons
        ttk.Button(self.export_frame, text="Export Chart", command=self.export_chart).grid(row=0, column=0, padx=5)
        ttk.Button(self.export_frame, text="Export Data", command=self.export_data).grid(row=0, column=1, padx=5)
        
        # Store results for export
        self.current_results = None
        self.current_er_date = None

    def create_column_controls(self):
        """Create checkboxes for column visibility control"""
        # Clear existing controls
        for widget in self.column_frame.winfo_children():
            widget.destroy()
            
        # Add label
        ttk.Label(self.column_frame, text="Show/Hide Columns:").grid(row=0, column=0, columnspan=4, sticky="w")
        
        # Create checkboxes in a grid layout
        row = 1
        col = 0
        for column in self.all_columns.keys():
            var = tk.BooleanVar(value=self.all_columns[column])
            cb = ttk.Checkbutton(
                self.column_frame,
                text=column,
                variable=var,
                command=lambda c=column, v=var: self.toggle_column(c, v)
            )
            cb.grid(row=row, column=col, sticky="w", padx=5, pady=2)
            col += 1
            if col > 3:  # 4 columns of checkboxes
                col = 0
                row += 1

    def toggle_column(self, column: str, var: tk.BooleanVar):
        """Toggle column visibility"""
        self.all_columns[column] = var.get()
        self.refresh_display()

    def refresh_display(self):
        """Refresh the display with current column settings"""
        if hasattr(self, 'current_results') and self.current_results:
            self.display_summary(self.current_results, self.current_er_date)

    def populate_earnings_dates(self, event=None):
        """Populate earnings dates when ticker is entered"""
        ticker = self.ticker_entry.get().strip().upper()
        if ticker:
            try:
                dates = self.analyzer.get_earnings_dates(ticker)
                if dates:
                    # Format dates for display
                    date_strings = [d.strftime('%Y-%m-%d') for d in dates]
                    self.er_date_combo['values'] = date_strings
                    self.er_date_combo.set(date_strings[0] if date_strings else '')
                else:
                    messagebox.showwarning("Warning", f"No earnings dates found for {ticker}")
                    self.er_date_combo['values'] = []
                    self.er_date_combo.set('')
            except Exception as e:
                messagebox.showerror("Error", f"Error fetching earnings dates: {str(e)}")
                self.er_date_combo['values'] = []
                self.er_date_combo.set('')

    def run_analysis(self):
        """Execute the analysis"""
        try:
            # Get inputs
            main_ticker = self.ticker_entry.get().upper()
            er_date = pd.to_datetime(self.er_date_var.get()).tz_localize(None)  # Make timezone-naive
            days = int(self.days_entry.get())
            peers = [t.strip().upper() for t in self.peers_entry.get().split(',')]
            
            # Validate inputs
            if not main_ticker or pd.isna(er_date):
                raise ValueError("Please enter ticker and select earnings date")
                
            # Analysis window
            start_date = er_date - timedelta(days=days)
            end_date = er_date + timedelta(days=days)
            
            # Get data
            results = {}
            results[main_ticker] = self.analyzer.get_stock_data(main_ticker, start_date, end_date)
            
            for peer in peers:
                if peer:
                    results[peer] = self.analyzer.get_stock_data(peer, start_date, end_date)
            
            # Store results for export
            self.current_results = results
            self.current_er_date = er_date
                    
            # Display results
            self.display_summary(results, er_date)
            self.display_charts(results, er_date)
            
            # Automatically export results
            self.export_chart()
            self.export_data()
            
        except Exception as e:
            messagebox.showerror("Error", str(e))
            
    def display_summary(self, results: Dict[str, pd.DataFrame], er_date: datetime):
        """Display summary statistics"""
        for widget in self.summary_frame.winfo_children():
            widget.destroy()
            
        # Get visible columns
        visible_columns = [col for col, shown in self.all_columns.items() if shown]
        
        # Create treeview with visible columns
        tree = ttk.Treeview(self.summary_frame, columns=visible_columns, show='headings')
        
        for col in visible_columns:
            tree.heading(col, text=col)
            tree.column(col, width=100)
            
        # Get main ticker data for correlation comparison
        main_ticker = list(results.keys())[0]
        main_data = results[main_ticker]['Close'] if results[main_ticker] is not None else None
            
        for ticker, data in results.items():
            if data is not None and not data.empty:
                # Get current values
                current_price = self.analyzer.get_current_price(ticker)
                current_iv = self.analyzer.get_current_iv(ticker)
                
                er_date_naive = pd.to_datetime(er_date).tz_localize(None)
                er_idx = data.index.searchsorted(er_date_naive)
                
                # Calculate correlation
                correlation = None
                if ticker != main_ticker:
                    correlation = self.analyzer.calculate_correlation(main_data, data['Close'])
                corr_category = self.analyzer.get_correlation_category(correlation) if ticker != main_ticker else "MAIN"
                
                if er_idx > 0 and er_idx < len(data):
                    # Get pre and post earnings dates
                    pre_er_date = data.index[er_idx - 1]
                    post_er_date = data.index[er_idx]
                    
                    # Get prices
                    pre_price = data['Close'].iloc[er_idx - 1]
                    post_price = data['Open'].iloc[er_idx]
                    price_change = ((post_price / pre_price) - 1) * 100
                    
                    # Get IVs
                    pre_iv = self.analyzer.get_historical_iv(ticker, pre_er_date)
                    post_iv = self.analyzer.get_historical_iv(ticker, post_er_date)
                    iv_change = ((post_iv / pre_iv) - 1) * 100 if (pre_iv and post_iv) else None
                    
                    # Calculate other metrics
                    pre_data = data.iloc[:er_idx]
                    post_data = data.iloc[er_idx:]
                    pre_return = pre_data['Daily_Return'].sum() if not pre_data.empty else None
                    post_return = post_data['Daily_Return'].sum() if not post_data.empty else None
                    vol_change = ((post_data['Volume'].mean() / pre_data['Volume'].mean()) - 1) * 100
                    
                    # Get latest technical indicators
                    latest = data.iloc[-1]
                    current_rsi = latest['RSI'] if 'RSI' in data.columns else None
                    price_vs_ma200 = (latest['Close'] / latest['MA200'] - 1) * 100 if ('MA200' in data.columns and pd.notnull(latest['MA200'])) else None
                    price_vs_ma50 = (latest['Close'] / latest['MA50'] - 1) * 100 if ('MA50' in data.columns and pd.notnull(latest['MA50'])) else None
                    
                    # Get MA signals
                    ma_signals = self.analyzer.check_ma_signals(data)
                    ma_cross_text = "50MA > 200MA" if ma_signals.get('golden_cross', False) else "50MA < 200MA"
                    
                    # Create values dictionary for all possible columns
                    values_dict = {
                        'Ticker': ticker,
                        'Current Price': f"${current_price:.2f}" if current_price is not None else "N/A",
                        'Pre-ER Price': f"${pre_price:.2f}" if pre_price is not None else "N/A",
                        'Post-ER Price': f"${post_price:.2f}" if post_price is not None else "N/A",
                        'Price Change': f"{price_change:+.2f}%" if price_change is not None else "N/A",
                        'Current IV': f"{current_iv:.1%}" if current_iv is not None else "N/A",
                        'Pre-ER IV': f"{pre_iv:.1%}" if pre_iv is not None else "N/A",
                        'Post-ER IV': f"{post_iv:.1%}" if post_iv is not None else "N/A",
                        'IV Change': f"{iv_change:+.1f}%" if iv_change is not None else "N/A",
                        'Pre-ER Return': f"{pre_return:.2%}" if pre_return is not None else "N/A",
                        'Post-ER Return': f"{post_return:.2%}" if post_return is not None else "N/A",
                        'Volume Change': f"{vol_change:.1f}%" if vol_change is not None else "N/A",
                        'Current RSI': f"{current_rsi:.1f}" if current_rsi is not None else "N/A",
                        'Price vs MA200': f"{price_vs_ma200:+.1f}%" if price_vs_ma200 is not None else "N/A",
                        'Price vs MA50': f"{price_vs_ma50:+.1f}%" if price_vs_ma50 is not None else "N/A",
                        'MA Cross': ma_cross_text,
                        'Correlation': corr_category
                    }
                    
                    # Insert only visible columns
                    tree.insert('', 'end', values=[values_dict[col] for col in visible_columns])
                else:
                    tree.insert('', 'end', values=(ticker, "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A"))
                    
        tree.grid(sticky="nsew")
        
    def display_charts(self, results: Dict[str, pd.DataFrame], er_date: datetime):
        """Display analysis charts"""
        for widget in self.charts_frame.winfo_children():
            widget.destroy()
            
        # Create figure with 4 subplots
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(15, 12))
        
        er_date_naive = pd.to_datetime(er_date).tz_localize(None)
        
        for ticker, data in results.items():
            if data is not None:
                # Returns plot
                ax1.plot(data.index, data['Cumulative_Return'], label=ticker)
                
                # Volume plot
                ax2.plot(data.index, data['Volume'], label=ticker)
                
                # RSI plot
                ax3.plot(data.index, data['RSI'], label=ticker)
                ax3.axhline(y=70, color='r', linestyle='--', alpha=0.5)
                ax3.axhline(y=30, color='g', linestyle='--', alpha=0.5)
                
                # Price and MAs plot
                ax4.plot(data.index, data['Close'], label=f'{ticker} Price')
                ax4.plot(data.index, data['MA50'], label=f'{ticker} MA50', linestyle='--', alpha=0.7)
                ax4.plot(data.index, data['MA200'], label=f'{ticker} MA200', linestyle='--', alpha=0.7)
                
        # Add vertical lines for earnings date
        for ax in [ax1, ax2, ax3, ax4]:
            ax.axvline(x=er_date_naive, color='r', linestyle='--')
            ax.legend()
            ax.grid(True)
        
        ax1.set_title('Returns Comparison')
        ax2.set_title('Volume Comparison')
        ax3.set_title('RSI (14-day)')
        ax4.set_title('Price and Moving Averages')
        
        plt.tight_layout()
        
        canvas = FigureCanvasTkAgg(fig, self.charts_frame)
        canvas.draw()
        canvas.get_tk_widget().grid(sticky="nsew")

    def export_chart(self):
        """Export the current chart as PNG"""
        if not hasattr(self, 'current_results') or not self.current_results:
            messagebox.showwarning("Warning", "No analysis results to export")
            return
            
        try:
            main_ticker = self.ticker_entry.get().upper()
            filename = f"earnings_analysis_{main_ticker}_{self.current_er_date.strftime('%Y%m%d')}.png"
            
            # Create a new figure for export (higher resolution)
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 10), dpi=300)
            
            for ticker, data in self.current_results.items():
                if data is not None:
                    ax1.plot(data.index, data['Cumulative_Return'], label=ticker)
                    ax2.plot(data.index, data['Volume'], label=ticker)
                    
            ax1.axvline(x=self.current_er_date, color='r', linestyle='--')
            ax2.axvline(x=self.current_er_date, color='r', linestyle='--')
            
            ax1.set_title('Returns Comparison')
            ax2.set_title('Volume Comparison')
            
            ax1.legend()
            ax2.legend()
            
            ax1.grid(True)
            ax2.grid(True)
            
            plt.tight_layout()
            plt.savefig(filename)
            plt.close(fig)
            
            messagebox.showinfo("Success", f"Chart exported as {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export chart: {str(e)}")

    def export_data(self):
        """Export the analysis data as CSV"""
        if not hasattr(self, 'current_results') or not self.current_results:
            messagebox.showwarning("Warning", "No analysis results to export")
            return
            
        try:
            main_ticker = self.ticker_entry.get().upper()
            filename = f"earnings_analysis_{main_ticker}_{self.current_er_date.strftime('%Y%m%d')}.csv"
            
            all_data = pd.DataFrame()
            
            for ticker, data in self.current_results.items():
                if data is not None:
                    data = data.copy()
                    # Add all relevant columns
                    data[f'{ticker}_Return'] = data['Daily_Return']
                    data[f'{ticker}_Cumulative'] = data['Cumulative_Return']
                    data[f'{ticker}_Volume'] = data['Volume']
                    data[f'{ticker}_RSI'] = data['RSI']
                    data[f'{ticker}_MA50'] = data['MA50']
                    data[f'{ticker}_MA200'] = data['MA200']
                    data[f'{ticker}_IV'] = data['IV']
                    
                    if all_data.empty:
                        all_data = data[[f'{ticker}_Return', f'{ticker}_Cumulative', f'{ticker}_Volume',
                                       f'{ticker}_RSI', f'{ticker}_MA50', f'{ticker}_MA200', f'{ticker}_IV']]
                    else:
                        all_data = all_data.join(data[[f'{ticker}_Return', f'{ticker}_Cumulative', f'{ticker}_Volume',
                                                      f'{ticker}_RSI', f'{ticker}_MA50', f'{ticker}_MA200', f'{ticker}_IV']])
            
            all_data.to_csv(filename)
            messagebox.showinfo("Success", f"Data exported as {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export data: {str(e)}")

    def export_report(self):
        """Create a Word document report of the analysis"""
        # ... existing code ...
        
        # Add correlation analysis
        doc = Document()
        main_ticker = list(self.current_results.keys())[0]
        main_data = self.current_results[main_ticker]['Close'] if self.current_results[main_ticker] is not None else None
        
        for ticker, data in self.current_results.items():
            if ticker != main_ticker and data is not None:
                correlation = self.analyzer.calculate_correlation(main_data, data['Close'])
                corr_category = self.analyzer.get_correlation_category(correlation)
                doc.add_paragraph(f'{ticker} correlation with {main_ticker}: {corr_category}', style='List Bullet')
        
        # ... rest of export_report code ...

    def export_setup_guide(self):
        """Create a Word document with setup instructions"""
        try:
            doc = Document()
            
            # Title
            doc.add_heading('Stock Earnings Analysis - Setup Guide', 0)
            
            # 1. Clone Repository
            doc.add_heading('1. Clone the Repository', level=1)
            p = doc.add_paragraph()
            p.add_run('Navigate to where you want to store the project:\n').bold = True
            doc.add_paragraph('cd c:\\   # or any preferred directory')
            doc.add_paragraph('git clone https://github.com/messfin/stock-earnings-analysis.git')
            doc.add_paragraph('cd stock-earnings-analysis')
            
            # 2. Set Up Python Environment
            doc.add_heading('2. Set Up Python Environment', level=1)
            doc.add_heading('Option A: Using venv (Recommended)', level=2)
            doc.add_paragraph('# Create virtual environment\npython -m venv venv')
            p = doc.add_paragraph('# Activate virtual environment\n')
            p.add_run('On Windows:\n').bold = True
            doc.add_paragraph('venv\\Scripts\\activate')
            p = doc.add_paragraph()
            p.add_run('On Mac/Linux:\n').bold = True
            doc.add_paragraph('source venv/bin/activate')
            
            doc.add_heading('Option B: Direct Installation', level=2)
            doc.add_paragraph('If you don\'t want to use a virtual environment, you can install packages directly.')
            
            # 3. Install Required Packages
            doc.add_heading('3. Install Required Packages', level=1)
            doc.add_paragraph('# Install all required packages\npip install -r requirements.txt')
            doc.add_paragraph('# If requirements.txt is missing, install packages directly:\npip install yfinance pandas matplotlib python-docx numpy tkinter')
            
            # 4. Run the Scripts
            doc.add_heading('4. Run the Scripts', level=1)
            doc.add_heading('To run the sector analysis:', level=2)
            doc.add_paragraph('python earnings_sector_compare.py')
            p = doc.add_paragraph('This will open the GUI application where you can:')
            doc.add_paragraph('- Enter a stock ticker', style='List Bullet')
            doc.add_paragraph('- Select earnings dates', style='List Bullet')
            doc.add_paragraph('- Add peer companies for comparison', style='List Bullet')
            doc.add_paragraph('- View and export analysis results', style='List Bullet')
            
            doc.add_heading('To run the main analysis:', level=2)
            doc.add_paragraph('python zmtec_hmain.py')
            
            # 5. Troubleshooting
            doc.add_heading('5. Troubleshooting', level=1)
            doc.add_paragraph('If you encounter any errors:')
            doc.add_paragraph('1. Verify Python installation:\npython --version  # Should show Python 3.x')
            doc.add_paragraph('2. Check installed packages:\npip list')
            
            p = doc.add_paragraph('3. Common issues:')
            doc.add_paragraph('- If you get "Module not found" errors, try:\n  pip install [module_name]', style='List Bullet')
            doc.add_paragraph('- If you get permission errors, try running as administrator', style='List Bullet')
            doc.add_paragraph('- For tkinter issues on Linux:\n  sudo apt-get install python3-tk  # For Ubuntu/Debian', style='List Bullet')
            
            # 6. Data Export Location
            doc.add_heading('6. Data Export Location', level=1)
            doc.add_paragraph('- Charts and data files will be saved in the same directory as the scripts', style='List Bullet')
            doc.add_paragraph('- Look for files with names like:', style='List Bullet')
            doc.add_paragraph('  - earnings_analysis_[TICKER]_[DATE].png', style='List Bullet')
            doc.add_paragraph('  - earnings_analysis_[TICKER]_[DATE].csv', style='List Bullet')
            
            # 7. Updates
            doc.add_heading('7. Updates', level=1)
            doc.add_paragraph('To get the latest updates:\ngit pull origin main')
            
            # 8. Example Usage
            doc.add_heading('8. Example Usage', level=1)
            p = doc.add_paragraph('For earnings_sector_compare.py:')
            doc.add_paragraph('1. Enter a ticker (e.g., "AAPL")', style='List Number')
            doc.add_paragraph('2. Select an earnings date from the dropdown', style='List Number')
            doc.add_paragraph('3. Enter peer tickers (e.g., "MSFT,GOOGL")', style='List Number')
            doc.add_paragraph('4. Click "Analyze"', style='List Number')
            
            p = doc.add_paragraph('\nFor zmtec_hmain.py:')
            doc.add_paragraph('1. Follow the command-line prompts', style='List Number')
            doc.add_paragraph('2. Enter requested information', style='List Number')
            doc.add_paragraph('3. View the analysis results', style='List Number')
            
            # Need Help section
            doc.add_heading('Need Help?', level=1)
            doc.add_paragraph('- Check the repository issues page', style='List Bullet')
            doc.add_paragraph('- Review the documentation in the code', style='List Bullet')
            doc.add_paragraph('- Contact the repository maintainer', style='List Bullet')
            
            # Save the document with a specific path
            filename = r'c:\sa\Stock_Earnings_Analysis_Setup_Guide.docx'  # Example path
            doc.save(filename)
            messagebox.showinfo("Success", f"Setup guide exported as {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export setup guide: {str(e)}")

    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = ERAnalysisApp()
    app.run()